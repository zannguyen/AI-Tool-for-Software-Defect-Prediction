"""
Create Week 7 Documentation
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Week 7: Final Report and Presentation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Introduction
doc.add_paragraph(
    'This week focuses on preparing final deliverables including the project report and presentation.'
)

# Final Report Structure
doc.add_heading('1. Final Project Report Structure', level=1)

sections = [
    '1. Introduction',
    '2. Related Work',
    '3. System Architecture',
    '4. Measurement Model',
    '5. AI Methods',
    '6. Implementation',
    '7. Experiments',
    '8. Results',
    '9. Conclusion',
    '10. References',
    '11. Appendices'
]

for section in sections:
    doc.add_paragraph(section)

# Section Details
doc.add_heading('2. Report Section Details', level=1)

# 1. Introduction
doc.add_heading('2.1 Introduction (1-2 pages)', level=2)

intro_content = [
    'Background: Software defects and their impact',
    'Problem Statement: Need for early defect detection',
    'Objectives: Build AI tool for defect prediction',
    'Scope: What the project covers and does not cover',
    'Overview: Brief summary of the approach'
]

for content in intro_content:
    doc.add_paragraph(content, style='List Bullet')

# 2. Related Work
doc.add_heading('2.2 Related Work (2-3 pages)', level=2)

related = [
    'Software defect prediction research',
    'Machine learning in software engineering',
    'Existing tools and approaches',
    'Comparison with our approach'
]

for r in related:
    doc.add_paragraph(r, style='List Bullet')

# 3. System Architecture
doc.add_heading('2.3 System Architecture (2-3 pages)', level=2)

arch = [
    'Overall system design',
    'Data flow diagram',
    'Component design',
    'Technology stack'
]

for a in arch:
    doc.add_paragraph(a, style='List Bullet')

# 4. Measurement Model
doc.add_heading('2.4 Measurement Model (2-3 pages)', level=2)

mm = [
    'Software metrics used',
    'Metrics calculation methods',
    'Supported languages',
    'Data preprocessing'
]

for m in mm:
    doc.add_paragraph(m, style='List Bullet')

# 5. AI Methods
doc.add_heading('2.5 AI Methods (2-3 pages)', level=2)

ai = [
    'Logistic Regression - theory and application',
    'Random Forest - theory and application',
    'Neural Network - theory and application',
    'Training process',
    'Ensemble methods'
]

for a in ai:
    doc.add_paragraph(a, style='List Bullet')

# 6. Implementation
doc.add_heading('2.6 Implementation (2-3 pages)', level=2)

impl = [
    'Code structure',
    'Key algorithms',
    'User interface',
    'Database design'
]

for i in impl:
    doc.add_paragraph(i, style='List Bullet')

# 7. Experiments
doc.add_heading('2.7 Experiments (2-3 pages)', level=2)

exp = [
    'Dataset description',
    'Experimental setup',
    'Evaluation metrics',
    'Comparison criteria'
]

for e in exp:
    doc.add_paragraph(e, style='List Bullet')

# 8. Results
doc.add_heading('2.8 Results (2-3 pages)', level=2)

results = [
    'Model performance comparison',
    'Analysis of results',
    'Discussion of findings',
    'Visualizations and charts'
]

for r in results:
    doc.add_paragraph(r, style='List Bullet')

# 9. Conclusion
doc.add_heading('2.9 Conclusion (1 page)', level=2)

conclusion = [
    'Summary of achievements',
    'Limitations',
    'Future work',
    'Lessons learned'
]

for c in conclusion:
    doc.add_paragraph(c, style='List Bullet')

# Presentation
doc.add_heading('3. Presentation Guidelines', level=1)

doc.add_paragraph('Presentation Duration: 15 minutes')

presentation = [
    'Introduction (2 min): Problem and objectives',
    'Related Work (2 min): Brief literature review',
    'System Design (3 min): Architecture and design',
    'Implementation (3 min): Demo and code walkthrough',
    'Experiments (3 min): Results and analysis',
    'Conclusion (2 min): Summary and future work'
]

for p in presentation:
    doc.add_paragraph(p, style='List Bullet')

# Presentation Tips
doc.add_heading('4. Presentation Tips', level=1)

tips = [
    'Practice multiple times before the actual presentation',
    'Use visuals and diagrams to explain concepts',
    'Focus on key results and insights',
    'Be prepared to answer questions',
    'Keep slides clean and readable',
    'Time management is crucial'
]

for tip in tips:
    doc.add_paragraph(tip, style='List Bullet')

# Software Demo
doc.add_heading('5. Software Demo', level=1)

demo = [
    'Show the web interface',
    'Demonstrate data upload',
    'Show metrics extraction',
    'Run model training',
    'Display predictions',
    'Show evaluation metrics',
    'Demonstrate report generation'
]

for d in demo:
    doc.add_paragraph(d, style='List Bullet')

# Deliverables Checklist
doc.add_heading('6. Deliverables Checklist', level=1)

deliverables = [
    'Final Project Report (15-20 pages)',
    'Presentation slides (15 minutes)',
    'Source code with documentation',
    'Trained models',
    'Sample datasets used',
    'User manual (optional)'
]

for d in deliverables:
    doc.add_paragraph(d, style='List Bullet')

# Submission Requirements
doc.add_heading('7. Submission Requirements', level=1)

submission = [
    'Submit all files via course portal',
    'Include README with running instructions',
    'Provide sample data for testing',
    'Include evaluation results'
]

for s in submission:
    doc.add_paragraph(s, style='List Bullet')

# Final Notes
doc.add_heading('8. Final Notes', level=1)

notes = [
    'Review and proofread the entire report',
    'Check all citations and references',
    'Ensure all figures and tables are properly labeled',
    'Test the software one final time',
    'Prepare for Q&A session',
    'Backup all files'
]

for note in notes:
    doc.add_paragraph(note, style='List Bullet')

# Save
doc.save('document/Week7_Final_Report_Presentation.docx')
print("Week 7 document created successfully!")
