"""
Create Week 2 Documentation
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Week 2: Requirements and System Design', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. System Requirements
doc.add_heading('1. System Requirements', level=1)

doc.add_paragraph(
    'The AI Tool for Software Defect Prediction requires the following functional and '
    'non-functional requirements:'
)

doc.add_heading('1.1 Functional Requirements', level=2)

fr = [
    'FR-01: The system shall accept CSV files containing software metrics',
    'FR-02: The system shall extract metrics from source code files (.py, .java, .js, .c, .cpp, .cs)',
    'FR-03: The system shall preprocess data (handle missing values, normalize)',
    'FR-04: The system shall train Logistic Regression model',
    'FR-05: The system shall train Random Forest model',
    'FR-06: The system shall train Neural Network (MLP) model',
    'FR-07: The system shall evaluate models using Accuracy, Precision, Recall, F1-Score, ROC-AUC',
    'FR-08: The system shall display prediction results with risk levels (High/Medium/Low)',
    'FR-09: The system shall generate visualizations (bar charts, heatmaps, ROC curves)',
    'FR-10: The system shall export reports (CSV, TXT)',
    'FR-11: The system shall store analysis history in database',
    'FR-12: The system shall save trained models for future use'
]

for req in fr:
    doc.add_paragraph(req)

doc.add_heading('1.2 Non-Functional Requirements', level=2)

nfr = [
    'NFR-01: Response time for prediction shall be less than 5 seconds',
    'NFR-02: System shall support dataset with at least 1000 samples',
    'NFR-03: User interface shall be intuitive and user-friendly',
    'NFR-04: System shall work on modern browsers (Chrome, Firefox, Edge)',
    'NFR-05: Code shall be well-documented and maintainable'
]

for req in nfr:
    doc.add_paragraph(req)

# 2. Input Artifacts
doc.add_heading('2. Input Artifacts', level=1)

doc.add_paragraph(
    'The system accepts the following input types:'
)

inputs = [
    'CSV Files: Software metrics in tabular format with LABEL column',
    'Source Code Files: Python, Java, JavaScript, C/C++, C# files',
    'Pre-processed Datasets: NASA KC1, KC2 datasets (ARFF format)'
]

for inp in inputs:
    doc.add_paragraph(inp, style='List Bullet')

# 3. System Architecture
doc.add_heading('3. System Architecture', level=1)

doc.add_paragraph(
    'The system follows a layered architecture:'
)

layers = [
    'Layer 1 - Presentation Layer: Streamlit web interface',
    'Layer 2 - Business Logic Layer: Data processing, model training, prediction',
    'Layer 3 - Data Layer: SQLite database for history, model files for persistence'
]

for layer in layers:
    doc.add_paragraph(layer, style='List Number')

# Data Flow
doc.add_heading('4. Data Flow', level=1)

doc.add_paragraph('Input → Processing → Metrics → Dashboard')

doc.add_heading('4.1 Input Stage', level=2)
doc.add_paragraph('Users upload CSV files or source code files through the web interface.')

doc.add_heading('4.2 Processing Stage', level=2)
doc.add_paragraph(
    'The system processes data through: '
    'Metrics Extraction → Data Cleaning → Feature Selection → Model Training'
)

doc.add_heading('4.3 Metrics Stage', level=2)
doc.add_paragraph(
    'Software metrics are calculated: LOC, Cyclomatic Complexity, Coupling, '
    'Function Count, Class Count, Comment Ratio, etc.'
)

doc.add_heading('4.4 Dashboard Stage', level=2)
doc.add_paragraph(
    'Results are displayed through interactive dashboards with charts, tables, and reports.'
)

# Technology Selection
doc.add_heading('5. Technology Selection', level=1)

tech = [
    'Programming Language: Python 3.8+',
    'Web Framework: Streamlit',
    'Machine Learning: Scikit-learn',
    'Deep Learning: TensorFlow/Keras (optional)',
    'Data Processing: Pandas, NumPy',
    'Visualization: Plotly, Matplotlib',
    'Database: SQLite',
    'Development Tools: VSCode, Git'
]

for t in tech:
    doc.add_paragraph(t, style='List Bullet')

# Dataset Identification
doc.add_heading('6. Dataset Identification', level=1)

doc.add_paragraph('The project uses the following datasets:')

datasets = [
    'NASA KC1 Dataset: 2,109 software modules with defects',
    'NASA KC2 Dataset: 522 software modules with defects',
    'Combined KC1+KC2: 2,631 samples with 16.5% defect rate',
    'Custom Dataset: User can upload their own CSV files'
]

for ds in datasets:
    doc.add_paragraph(ds, style='List Bullet')

# Save
doc.save('document/Week2_Requirements_System_Design.docx')
print("Week 2 document created successfully!")
