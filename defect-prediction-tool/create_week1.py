"""
Create Week 1 Documentation
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Create document
doc = Document()

# Title
title = doc.add_heading('Week 1: Project Selection and Proposal', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Project Title
doc.add_heading('Project Title', level=1)
p = doc.add_paragraph('P7 - AI Tool for Software Defect Prediction')
p.runs[0].bold = True

# Problem Statement
doc.add_heading('1. Problem Statement', level=1)
doc.add_paragraph(
    'Software defects are a major concern in software engineering, causing significant '
    'cost and time overruns. Early detection of potential defects can save substantial '
    'resources and improve software quality. This project aims to develop an AI-powered '
    'tool that predicts software defects using machine learning techniques based on '
    'code metrics.'
)

# Measurement Model
doc.add_heading('2. Measurement Model Used', level=1)
doc.add_paragraph(
    'The project uses software metrics as the measurement model. These metrics include:'
)

metrics = [
    'LOC (Lines of Code) - Total lines of code in a module',
    'Cyclomatic Complexity - Measure of program complexity based on decision points',
    'Coupling - Degree of interdependence between modules',
    'Code Churn - Number of lines changed over time',
    'Function Count - Number of functions/methods in a module',
    'Class Count - Number of classes in a module',
    'Comment Ratio - Ratio of comment lines to total lines',
    'Decision Count - Number of decision points in code',
    'Essential Complexity - Complexity due to unstructured code',
    'Halstead Metrics - Volume, difficulty, effort, and estimated bugs'
]

for metric in metrics:
    doc.add_paragraph(metric, style='List Bullet')

# AI Technique
doc.add_heading('3. AI Technique Used', level=1)
doc.add_paragraph(
    'The project employs three machine learning algorithms for defect prediction:'
)

ai_techniques = [
    'Logistic Regression - Linear model for binary classification',
    'Random Forest - Ensemble learning method using decision trees',
    'Neural Network (MLP) - Multi-layer perceptron for complex pattern recognition'
]

for tech in ai_techniques:
    doc.add_paragraph(tech, style='List Bullet')

doc.add_paragraph(
    'Note: TensorFlow/Keras will be used if available; otherwise, sklearn MLPClassifier '
    'will serve as the fallback implementation.'
)

# System Architecture
doc.add_heading('4. System Architecture', level=1)

doc.add_paragraph('The system follows a modular architecture:')

arch_components = [
    '1. Data Input Layer - Accepts CSV files or source code files',
    '2. Metrics Extraction Layer - Extracts software metrics from code',
    '3. Preprocessing Layer - Cleans and transforms data',
    '4. ML Training Layer - Trains prediction models',
    '5. Prediction Layer - Generates defect predictions',
    '6. Evaluation Layer - Calculates performance metrics',
    '7. Dashboard Layer - Visualizes results',
    '8. Report Generation Layer - Exports analysis reports'
]

for comp in arch_components:
    doc.add_paragraph(comp, style='List Number')

# Expected Outputs
doc.add_heading('5. Expected Outputs', level=1)
doc.add_paragraph('The project will deliver:')

outputs = [
    'A functional web-based tool for software defect prediction',
    'A dashboard displaying risk heatmaps and prediction results',
    'Evaluation metrics comparing different ML models',
    'A comprehensive project report documenting the development process',
    'Trained ML models saved for future use'
]

for output in outputs:
    doc.add_paragraph(output, style='List Bullet')

# Additional Information
doc.add_heading('Additional Information', level=1)

doc.add_paragraph('Course: Software Measurement & Analysis')
doc.add_paragraph('Project Topic: P7 - AI Tool for Software Defect Prediction')
doc.add_paragraph('Team Size: 5 students')

# Save
doc.save('document/Week1_Project_Selection_Proposal.docx')
print("Week 1 document created successfully!")
