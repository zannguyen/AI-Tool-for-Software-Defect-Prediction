"""
Create Week 3 Documentation
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Week 3: Measurement Model Implementation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Introduction
doc.add_paragraph(
    'This week focuses on implementing the core software measurement model that serves as the '
    'foundation for defect prediction. The measurement model calculates various software metrics '
    'that are used as features for the machine learning models.'
)

# Measurement Model Overview
doc.add_heading('1. Measurement Model Overview', level=1)

doc.add_paragraph(
    'The measurement model calculates software metrics from source code. These metrics are '
    'categorized into:'
)

categories = [
    'Size Metrics: LOC, LOC_BLANK, LOC_COMMENTS, LOC_CODE',
    'Complexity Metrics: Cyclomatic Complexity, Essential Complexity, Design Complexity',
    'Structure Metrics: Function Count, Class Count, Decision Count',
    'Quality Metrics: Comment Ratio, Code Churn',
    'Halstead Metrics: Volume, Difficulty, Effort, Estimated Bugs'
]

for cat in categories:
    doc.add_paragraph(cat, style='List Bullet')

# Metrics Calculation
doc.add_heading('2. Metrics Calculation Implementation', level=1)

doc.add_heading('2.1 Size Metrics', level=2)

size_metrics = [
    'LOC (Lines of Code): Total number of lines in the file',
    'LOC_BLANK: Number of blank lines',
    'LOC_COMMENTS: Number of comment lines',
    'LOC_CODE: Number of lines containing code',
    'LOC_TOTAL: Total lines including blank and comments'
]

for m in size_metrics:
    doc.add_paragraph(m)

doc.add_heading('2.2 Complexity Metrics', level=2)

complex_metrics = [
    'CYCLOMATIC_COMPLEXITY: Measures the number of linearly independent paths through code',
    'ESSENTIAL_COMPLEXITY: Complexity due to unstructured code',
    'DESIGN_COMPLEXITY: Complexity of the module design',
    'DECISION_COUNT: Number of decision points (if, while, for, switch, etc.)'
]

for m in complex_metrics:
    doc.add_paragraph(m)

doc.add_heading('2.3 Structure Metrics', level=2)

struct_metrics = [
    'FUNCTION_COUNT: Number of functions/methods defined in the code',
    'CLASS_COUNT: Number of classes defined in the code',
    'BRANCH_COUNT: Number of branch statements'
]

for m in struct_metrics:
    doc.add_paragraph(m)

doc.add_heading('2.4 Quality Metrics', level=2)

quality_metrics = [
    'COMMENT_RATIO: Ratio of comment lines to total lines (LOC_COMMENTS / LOC)',
    'CODE_CHURN: Number of lines changed over time (for prediction only)'
]

for m in quality_metrics:
    doc.add_paragraph(m)

# Halstead Metrics
doc.add_heading('2.5 Halstead Metrics', level=2)

doc.add_paragraph('These metrics are calculated from the NASA KC1/KC2 datasets:')

halstead = [
    'n: Program length - total number of operators and operands',
    'v: Program volume - information content of the program',
    'l: Program level - inverse of difficulty',
    'd: Program difficulty',
    'i: Program intelligence',
    'e: Program effort to understand',
    'b: Estimated number of bugs',
    't: Estimated time to understand (in seconds)',
    'uniq_Op: Number of unique operators',
    'uniq_Opnd: Number of unique operands',
    'total_Op: Total number of operators',
    'total_Opnd: Total number of operands'
]

for h in halstead:
    doc.add_paragraph(h)

# Supported Languages
doc.add_heading('3. Supported Programming Languages', level=1)

languages = [
    'Python (.py)',
    'Java (.java)',
    'JavaScript (.js)',
    'C (.c)',
    'C++ (.cpp, .h, .hpp)',
    'C# (.cs)'
]

for lang in languages:
    doc.add_paragraph(lang, style='List Bullet')

# Example Output
doc.add_heading('4. Example Output', level=1)

doc.add_paragraph('The measurement model produces the following output for each file:')

output_example = '''
File: example.py
LOC: 150
LOC_BLANK: 20
LOC_COMMENTS: 15
LOC_CODE: 115
FUNCTION_COUNT: 5
CLASS_COUNT: 2
CYCLOMATIC_COMPLEXITY: 12
DECISION_COUNT: 8
COMMENT_RATIO: 0.10
'''

doc.add_paragraph(output_example)

# Implementation Details
doc.add_heading('5. Implementation Details', level=1)

doc.add_paragraph(
    'The measurement model is implemented in the file: src/code_metrics_extractor.py'
)

doc.add_heading('5.1 Key Functions', level=2)

functions = [
    'extract_from_directory(directory_path): Extract metrics from all files in a directory',
    'extract_from_files(file_paths): Extract metrics from a list of file paths',
    'extract_python_metrics(content, filename): Extract metrics from Python code',
    'extract_java_metrics(content, filename): Extract metrics from Java code',
    'extract_js_metrics(content, filename): Extract metrics from JavaScript code',
    'extract_c_metrics(content, filename): Extract metrics from C code',
    'extract_cpp_metrics(content, filename): Extract metrics from C++ code',
    'extract_csharp_metrics(content, filename): Extract metrics from C# code'
]

for func in functions:
    doc.add_paragraph(func, style='List Bullet')

# Deliverable
doc.add_heading('6. Deliverable - Prototype Module', level=1)

doc.add_paragraph('The prototype module includes:')

deliverables = [
    'Functional metrics extraction from source code',
    'Support for 6 programming languages',
    'Output in DataFrame format compatible with ML models',
    'CSV export capability'
]

for d in deliverables:
    doc.add_paragraph(d, style='List Bullet')

# Save
doc.save('document/Week3_Measurement_Model_Implementation.docx')
print("Week 3 document created successfully!")
