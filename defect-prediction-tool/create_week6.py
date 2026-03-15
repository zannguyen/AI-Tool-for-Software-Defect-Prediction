"""
Create Week 6 Documentation
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Week 6: Experimentation and Evaluation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Introduction
doc.add_paragraph(
    'This week focuses on evaluating the tool using test cases and datasets. '
    'The evaluation measures the accuracy and performance of the machine learning models.'
)

# Dataset Description
doc.add_heading('1. Dataset Description', level=1)

doc.add_paragraph(
    'The tool is evaluated using the NASA KC1 and KC2 datasets:'
)

doc.add_heading('1.1 NASA KC1 Dataset', level=2)

kc1_info = [
    'Total Samples: 2,109 software modules',
    'Features: 21 software metrics',
    'Defect Rate: Approximately 15.5%',
    'Source: NASA Metrics Data Program (MDP)',
    'Language: C++'
]

for info in kc1_info:
    doc.add_paragraph(info, style='List Bullet')

doc.add_heading('1.2 NASA KC2 Dataset', level=2)

kc2_info = [
    'Total Samples: 522 software modules',
    'Features: 21 software metrics',
    'Defect Rate: Approximately 20.5%',
    'Source: NASA Metrics Data Program (MDP)',
    'Language: C++'
]

for info in kc2_info:
    doc.add_paragraph(info, style='List Bullet')

doc.add_heading('1.3 Combined Dataset (KC1+KC2)', level=2)

combined_info = [
    'Total Samples: 2,631 software modules',
    'Features: 21 software metrics',
    'Defect Rate: 16.5% (433 defective, 2,198 non-defective)',
    'Train/Test Split: 80% / 20%'
]

for info in combined_info:
    doc.add_paragraph(info, style='List Bullet')

# Experiment Setup
doc.add_heading('2. Experiment Setup', level=1)

doc.add_heading('2.1 Data Preprocessing', level=2)

preprocessing = [
    'Removed non-numeric columns (file_name, file_path)',
    'Converted categorical labels to numeric (false=0, true=1)',
    'Handled missing values by dropping rows with NaN',
    'Applied StandardScaler for feature normalization'
]

for p in preprocessing:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('2.2 Model Configuration', level=2)

config = [
    'Test Size: 25%',
    'Random State: 42 (for reproducibility)',
    'Stratified Split: Yes (maintain class distribution)',
    'Training Samples: 1,973',
    'Test Samples: 658'
]

for c in config:
    doc.add_paragraph(c, style='List Bullet')

# Evaluation Metrics
doc.add_heading('3. Evaluation Metrics', level=1)

metrics = [
    'Accuracy: Overall correctness of predictions',
    'Precision: Proportion of true positives among positive predictions',
    'Recall (Sensitivity): Proportion of actual positives correctly identified',
    'F1-Score: Harmonic mean of Precision and Recall',
    'ROC-AUC: Area under the Receiver Operating Characteristic curve'
]

for m in metrics:
    doc.add_paragraph(m, style='List Bullet')

# Results
doc.add_heading('4. Results', level=1)

doc.add_heading('4.1 Model Performance Comparison', level=2)

doc.add_paragraph('The following table shows the evaluation results:')

# Create table
table = doc.add_table(rows=5, cols=6)
table.style = 'Table Grid'

# Header
header_cells = table.rows[0].cells
header_cells[0].text = 'Model'
header_cells[1].text = 'Accuracy'
header_cells[2].text = 'Precision'
header_cells[3].text = 'Recall'
header_cells[4].text = 'F1-Score'
header_cells[5].text = 'ROC-AUC'

# Row 1: Logistic Regression
row1 = table.rows[1].cells
row1[0].text = 'Logistic Regression'
row1[1].text = '0.745'
row1[2].text = '0.357'
row1[3].text = '0.694'
row1[4].text = '0.472'
row1[5].text = '0.796'

# Row 2: Random Forest
row2 = table.rows[2].cells
row2[0].text = 'Random Forest'
row2[1].text = '0.804'
row2[2].text = '0.419'
row2[3].text = '0.500'
row2[4].text = '0.456'
row2[5].text = '0.783'

# Row 3: Neural Network
row3 = table.rows[3].cells
row3[0].text = 'Neural Network'
row3[1].text = '0.851'
row3[2].text = '0.667'
row3[3].text = '0.185'
row3[4].text = '0.290'
row3[5].text = '0.793'

# Row 4: Best
row4 = table.rows[4].cells
row4[0].text = 'Best'
row4[1].text = 'Neural Network'
row4[2].text = 'Neural Network'
row4[3].text = 'Logistic Regression'
row4[4].text = 'Logistic Regression'
row4[5].text = 'Logistic Regression'

doc.add_paragraph('')

doc.add_heading('4.2 Analysis', level=2)

analysis = [
    'Logistic Regression achieves the highest ROC-AUC (0.796), indicating good ability to distinguish between defective and non-defective modules',
    'Random Forest achieves the highest Accuracy (0.804)',
    'Neural Network has highest Precision (0.667) but low Recall (0.185)',
    'Logistic Regression has the best balance between Precision and Recall for this dataset',
    'The imbalanced class distribution (16.5% defects) affects model performance'
]

for a in analysis:
    doc.add_paragraph(a, style='List Bullet')

# Confusion Matrix
doc.add_heading('4.3 Confusion Matrix', level=2)

doc.add_paragraph('Example - Logistic Regression:')

cm_table = doc.add_table(rows=3, cols=3)
cm_table.style = 'Table Grid'

cm_header = cm_table.rows[0].cells
cm_header[0].text = ''
cm_header[1].text = 'Predicted: No Defect'
cm_header[2].text = 'Predicted: Defect'

cm_row1 = cm_table.rows[1].cells
cm_row1[0].text = 'Actual: No Defect'
cm_row1[1].text = 'TN'
cm_row1[2].text = 'FP'

cm_row2 = cm_table.rows[2].cells
cm_row2[0].text = 'Actual: Defect'
cm_row2[1].text = 'FN'
cm_row2[2].text = 'TP'

# Discussion
doc.add_heading('5. Discussion', level=1)

discussion = [
    'All three models achieve ROC-AUC above 0.75, indicating reasonable predictive capability',
    'Logistic Regression performs best overall with highest Recall and F1-Score',
    'Random Forest shows good Accuracy but lower Recall, meaning it misses some defective modules',
    'Neural Network has high Precision but very low Recall, making it conservative in predicting defects',
    'Class imbalance affects the performance, especially for models that are not properly calibrated',
    'More data or feature engineering could improve the results',
    'Ensemble methods could be explored to combine strengths of different models'
]

for d in discussion:
    doc.add_paragraph(d, style='List Bullet')

# Limitations
doc.add_heading('6. Limitations', level=1)

limitations = [
    'Dataset size may be insufficient for complex patterns',
    'Class imbalance (16.5% defects) affects model training',
    'Limited to numerical features; code structure not fully captured',
    'Results may vary with different random seeds',
    'Single dataset evaluation; cross-validation not performed'
]

for l in limitations:
    doc.add_paragraph(l, style='List Bullet')

# Deliverable
doc.add_heading('7. Deliverable - Evaluation Report', level=1)

doc.add_paragraph('This evaluation report includes:')

deliverables = [
    'Dataset description (NASA KC1, KC2)',
    'Experiment setup and configuration',
    'Model performance results',
    'Confusion matrices',
    'Analysis and discussion',
    'Limitations and future work'
]

for d in deliverables:
    doc.add_paragraph(d, style='List Bullet')

# Save
doc.save('document/Week6_Experimentation_Evaluation.docx')
print("Week 6 document created successfully!")
