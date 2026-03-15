"""
Create Week 5 Documentation
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Week 5: Tool Development and Dashboard', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Introduction
doc.add_paragraph(
    'This week focuses on building the complete tool interface with visualization dashboard. '
    'The tool is developed using Streamlit framework for rapid prototyping of data science applications.'
)

# Technology Stack
doc.add_heading('1. Technology Stack', level=1)

tech = [
    'Frontend: Streamlit (Python web framework)',
    'Backend: Python 3.x',
    'Machine Learning: Scikit-learn',
    'Database: SQLite',
    'Visualization: Plotly, Matplotlib',
    'Deployment: Local development server'
]

for t in tech:
    doc.add_paragraph(t, style='List Bullet')

# Project Structure
doc.add_heading('2. Project Structure', level=1)

structure = [
    'app.py - Main Streamlit application',
    'src/preprocessing.py - Data preprocessing',
    'src/models.py - ML models',
    'src/code_metrics_extractor.py - Metrics extraction',
    'src/database.py - SQLite database handler',
    'src/evaluation.py - Model evaluation',
    'dashboard/Dashboard.py - Dashboard visualization',
    'models/ - Saved trained models',
    'reports/ - Generated reports',
    'data/ - Dataset storage'
]

for s in structure:
    doc.add_paragraph(s, style='List Bullet')

# Features
doc.add_heading('3. Features', level=1)

# Feature 1: Data Upload
doc.add_heading('3.1 Data Upload', level=2)

doc.add_paragraph(
    'Users can upload data in two formats:'
)

upload_features = [
    'CSV File Upload: Accept CSV files with software metrics and LABEL column',
    'Source Code Upload: Accept multiple source code files (.py, .java, .js, etc.)',
    'Sample Data: Use built-in NASA KC1/KC2 dataset'
]

for f in upload_features:
    doc.add_paragraph(f, style='List Bullet')

# Feature 2: Metrics Calculation
doc.add_heading('3.2 Metrics Calculation', level=2)

metrics_features = [
    'Extract LOC (Lines of Code)',
    'Extract Cyclomatic Complexity',
    'Extract Function Count',
    'Extract Class Count',
    'Extract Comment Ratio',
    'Extract Decision Count',
    'Support for 6 programming languages'
]

for f in metrics_features:
    doc.add_paragraph(f, style='List Bullet')

# Feature 3: Model Training
doc.add_heading('3.3 Model Training', level=2)

training_features = [
    'Train Logistic Regression model',
    'Train Random Forest model',
    'Train Neural Network (MLP) model',
    'Configure test size (10-40%)',
    'Configure random state for reproducibility',
    'Option for stratified split'
]

for f in training_features:
    doc.add_paragraph(f, style='List Bullet')

# Feature 4: Prediction
doc.add_heading('3.4 Prediction', level=2)

prediction_features = [
    'Predict defects using test data',
    'Manual input for custom metrics',
    'Display probability for each model',
    'Ensemble prediction (average of all models)',
    'Risk level classification (High/Medium/Low)'
]

for f in prediction_features:
    doc.add_paragraph(f, style='List Bullet')

# Feature 5: Dashboard
doc.add_heading('3.5 Dashboard', level=2)

dashboard_features = [
    'Overview metrics (total samples, features, model status)',
    'Data distribution pie chart',
    'Model comparison radar chart',
    'Recent activity timeline'
]

for f in dashboard_features:
    doc.add_paragraph(f, style='List Bullet')

# Feature 6: Evaluation
doc.add_heading('3.6 Evaluation', level=2)

eval_features = [
    'Accuracy metric',
    'Precision metric',
    'Recall metric',
    'F1-Score metric',
    'ROC-AUC metric',
    'Confusion Matrix heatmap',
    'ROC Curve visualization',
    'Feature Importance chart'
]

for f in eval_features:
    doc.add_paragraph(f, style='List Bullet')

# Feature 7: Reports
doc.add_heading('3.7 Reports', level=2)

report_features = [
    'Export evaluation results to CSV',
    'Export prediction results to CSV',
    'Generate detailed text report',
    'Save trained models'
]

for f in report_features:
    doc.add_paragraph(f, style='List Bullet')

# Feature 8: History
doc.add_heading('3.8 History', level=2)

history_features = [
    'Store analysis sessions in SQLite database',
    'Display recent sessions in sidebar',
    'View session details with metrics',
    'View historical charts'
]

for f in history_features:
    doc.add_paragraph(f, style='List Bullet')

# User Interface
doc.add_heading('4. User Interface', level=1)

doc.add_paragraph('The tool provides a sidebar navigation with the following pages:')

pages = [
    'Home - Welcome screen with instructions',
    'Dashboard - Overview with visualizations',
    'Data Upload - Upload CSV or source code',
    'Model Training - Train ML models',
    'Predictions - View prediction results',
    'Evaluation - View model performance metrics',
    'Export Reports - Download reports',
    'History - View analysis history',
    'Models - Save/load trained models',
    'About - Project information'
]

for p in pages:
    doc.add_paragraph(p, style='List Bullet')

# Dashboard Screenshots Description
doc.add_heading('5. Dashboard Visualizations', level=1)

doc.add_paragraph('The dashboard includes:')

visualizations = [
    'Pie Chart: Data distribution (defect vs non-defect)',
    'Bar Chart: Model comparison',
    'Radar Chart: Multi-metric model comparison',
    'Heatmap: Confusion Matrix',
    'Line Chart: ROC Curve',
    'Horizontal Bar: Feature Importance'
]

for v in visualizations:
    doc.add_paragraph(v, style='List Bullet')

# Deliverable
doc.add_heading('6. Deliverable - Functional Software Prototype', level=1)

doc.add_paragraph('The prototype includes:')

deliverables = [
    'Fully functional web interface',
    'Data upload capability',
    'Metrics extraction from source code',
    'Three ML models for prediction',
    'Interactive dashboard with visualizations',
    'Evaluation metrics display',
    'Report generation and export',
    'History tracking with SQLite',
    'Model saving and loading'
]

for d in deliverables:
    doc.add_paragraph(d, style='List Bullet')

# Running the Tool
doc.add_heading('7. Running the Tool', level=1)

doc.add_paragraph('To run the tool:')

run_steps = [
    '1. Install dependencies: pip install -r requirements.txt',
    '2. Run the application: streamlit run app.py',
    '3. Open browser at: http://localhost:8501',
    '4. Navigate using the sidebar menu'
]

for step in run_steps:
    doc.add_paragraph(step)

# Save
doc.save('document/Week5_Tool_Development_Dashboard.docx')
print("Week 5 document created successfully!")
