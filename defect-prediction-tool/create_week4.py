"""
Create Week 4 Documentation
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Week 4: AI / Machine Learning Integration', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Introduction
doc.add_paragraph(
    'This week focuses on integrating AI and Machine Learning components into the system. '
    'The tool uses three ML algorithms to predict software defects based on the metrics '
    'extracted in Week 3.'
)

# ML Models Overview
doc.add_heading('1. Machine Learning Models', level=1)

doc.add_paragraph(
    'The system implements three machine learning models for defect prediction:'
)

# Model 1: Logistic Regression
doc.add_heading('1.1 Logistic Regression', level=2)

doc.add_paragraph(
    'Logistic Regression is a linear classification algorithm that predicts the probability '
    'of a binary outcome. It is simple, interpretable, and serves as a baseline model.'
)

lr_params = [
    'random_state: 42 (for reproducibility)',
    'max_iter: 1000 (maximum iterations)',
    'class_weight: balanced (to handle class imbalance)'
]

doc.add_paragraph('Key parameters:')
for p in lr_params:
    doc.add_paragraph(p, style='List Bullet')

# Model 2: Random Forest
doc.add_heading('1.2 Random Forest', level=2)

doc.add_paragraph(
    'Random Forest is an ensemble learning method that constructs multiple decision trees '
    'and outputs the mode of the classes. It handles non-linear relationships well.'
)

rf_params = [
    'n_estimators: 100 (number of trees)',
    'max_depth: 10 (maximum tree depth)',
    'min_samples_split: 5 (minimum samples to split)',
    'min_samples_leaf: 2 (minimum samples in leaf)',
    'class_weight: balanced (to handle class imbalance)'
]

doc.add_paragraph('Key parameters:')
for p in rf_params:
    doc.add_paragraph(p, style='List Bullet')

# Model 3: Neural Network
doc.add_heading('1.3 Neural Network (MLP)', level=2)

doc.add_paragraph(
    'Multi-Layer Perceptron (MLP) is a feedforward neural network that can learn '
    'complex non-linear patterns. It uses sklearn MLPClassifier as the implementation.'
)

nn_params = [
    'hidden_layer_sizes: (64, 32, 16) - three hidden layers',
    'activation: relu - Rectified Linear Unit activation',
    'solver: adam - Adam optimizer',
    'max_iter: 500 - maximum iterations',
    'early_stopping: True - stop if no improvement',
    'validation_fraction: 0.1 - 10% for validation'
]

doc.add_paragraph('Key parameters:')
for p in nn_params:
    doc.add_paragraph(p, style='List Bullet')

doc.add_paragraph(
    'Note: If TensorFlow/Keras is available, it will be used instead of sklearn MLPClassifier.'
)

# Training Process
doc.add_heading('2. Training Process', level=1)

training_steps = [
    '1. Data Splitting: Split data into training (80%) and testing (20%) sets',
    '2. Feature Scaling: Apply StandardScaler to normalize features',
    '3. Model Training: Train each model on the training set',
    '4. Prediction: Generate predictions on the test set',
    '5. Evaluation: Calculate performance metrics'
]

for step in training_steps:
    doc.add_paragraph(step)

# Feature Engineering
doc.add_heading('3. Feature Engineering', level=1)

doc.add_paragraph(
    'The system uses the following features extracted from software metrics:'
)

features = [
    'Size Features: LOC, LOC_BLANK, LOC_COMMENTS, LOC_CODE',
    'Complexity Features: CYCLOMATIC_COMPLEXITY, ESSENTIAL_COMPLEXITY, DESIGN_COMPLEXITY',
    'Structure Features: FUNCTION_COUNT, CLASS_COUNT, BRANCH_COUNT',
    'Quality Features: COMMENT_RATIO',
    'Halstead Features: n, v, l, d, i, e, b, t',
    'Operator/Operand Features: uniq_Op, uniq_Opnd, total_Op, total_Opnd'
]

for f in features:
    doc.add_paragraph(f, style='List Bullet')

# Data Preprocessing
doc.add_heading('4. Data Preprocessing', level=1)

preprocessing = [
    'Handle Missing Values: Fill with mean or drop rows',
    'Feature Selection: Remove non-numeric columns (file_name, file_path)',
    'Normalization: Apply StandardScaler to all features',
    'Class Balancing: Use class_weight=balanced for imbalanced data',
    'Train-Test Split: 80-20 split with stratification'
]

for p in preprocessing:
    doc.add_paragraph(p, style='List Bullet')

# Prediction Process
doc.add_heading('5. Prediction Process', level=1)

doc.add_paragraph('For new input:')

prediction_steps = [
    '1. Extract metrics from the new code/module',
    '2. Preprocess: Scale features using the trained scaler',
    '3. Predict: Get probability predictions from all three models',
    '4. Ensemble: Average probabilities from all models',
    '5. Classify: If probability >= 0.5, predict defect; otherwise no defect'
]

for step in prediction_steps:
    doc.add_paragraph(step)

# Risk Assessment
doc.add_heading('6. Risk Assessment', level=1)

doc.add_paragraph('Predictions are classified into risk levels:')

risk_levels = [
    'HIGH (Red): Defect probability >= 0.5',
    'MEDIUM (Orange): Defect probability >= 0.3 and < 0.5',
    'LOW (Green): Defect probability < 0.3'
]

for r in risk_levels:
    doc.add_paragraph(r, style='List Bullet')

# Deliverable
doc.add_heading('7. Deliverable - AI Module Prototype', level=1)

doc.add_paragraph('The AI module includes:')

deliverables = [
    'Three trained ML models: Logistic Regression, Random Forest, Neural Network',
    'Feature preprocessing pipeline',
    'Prediction API for new modules',
    'Risk level classification',
    'Ensemble prediction combining all three models'
]

for d in deliverables:
    doc.add_paragraph(d, style='List Bullet')

# Save
doc.save('document/Week4_AI_ML_Integration.docx')
print("Week 4 document created successfully!")
